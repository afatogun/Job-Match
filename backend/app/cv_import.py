"""Step 12 - turn an uploaded master CV into a structured profile."""

import io
import logging

from .ai import complete_structured
from .models import Profile

log = logging.getLogger(__name__)

MAX_CV_CHARS = 20000

SYSTEM = """You extract a candidate's CV into structured data.

Rules:
- Use ONLY what the CV actually says. Never invent employers, dates, metrics or skills.
- Leave a field empty rather than guessing.
- Keep achievement bullets close to the candidate's own wording.
- Dates stay in whatever format the CV uses.
"""


def extract_text(filename: str, data: bytes) -> str:
    """Plain text from a PDF or DOCX upload."""
    name = filename.lower()

    if name.endswith(".pdf"):
        from pypdf import PdfReader

        reader = PdfReader(io.BytesIO(data))
        pages = [(page.extract_text() or "") for page in reader.pages]
        return "\n".join(pages).strip()

    if name.endswith(".docx"):
        import docx

        document = docx.Document(io.BytesIO(data))
        parts = [p.text for p in document.paragraphs]
        for table in document.tables:
            for row in table.rows:
                parts.extend(cell.text for cell in row.cells)
        return "\n".join(parts).strip()

    if name.endswith(".txt"):
        return data.decode("utf-8", errors="replace").strip()

    raise ValueError("Unsupported file type. Upload a PDF, DOCX or TXT file.")


def structure_cv(text: str, model: str, filename: str = "") -> Profile:
    """CV text -> Profile via the model. Raises if the API is unusable."""
    if not text.strip():
        raise ValueError(
            "No text could be read from that file. If it is a scanned PDF, "
            "the text layer is missing - export a text-based PDF or upload DOCX."
        )

    profile = complete_structured(
        model=model,
        system=SYSTEM,
        user=f"Extract this CV into the structured format:\n\n{text[:MAX_CV_CHARS]}",
        schema_model=Profile,
    )
    profile.source_filename = filename
    return profile
