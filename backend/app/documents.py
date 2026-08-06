"""Step 17 - deterministic document rendering.

The model supplies content only. Layout lives here, so every CV comes out of the
same template: one column, no tables for layout, no text boxes, no icons, no skill
bars - all things that break ATS parsers.
"""

import logging
import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor

from .config import GENERATED_DIR
from .models import GeneratedCV

log = logging.getLogger(__name__)

ACCENT = RGBColor(0x1F, 0x2A, 0x37)
MUTED = RGBColor(0x55, 0x5F, 0x6D)


def slugify(value: str | None, fallback: str = "item") -> str:
    slug = re.sub(r"[^a-zA-Z0-9]+", "-", (value or "").strip().lower()).strip("-")
    return (slug or fallback)[:60]


def application_folder(job: dict) -> Path:
    """generated/{company}-{role}-{job_id}/"""
    company = str(job.get("company") or "")
    title = str(job.get("title") or "")
    name = f"{slugify(company, 'company')}-{slugify(title, 'role')}-{job['id']}"
    folder = GENERATED_DIR / name
    folder.mkdir(parents=True, exist_ok=True)
    return folder


def _configure_styles(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)
    pf = normal.paragraph_format
    pf.space_after = Pt(4)
    pf.line_spacing = 1.06

    for section in doc.sections:
        section.top_margin = section.bottom_margin = Pt(40)
        section.left_margin = section.right_margin = Pt(48)


def _heading(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(11)
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(text.upper())
    run.bold = True
    run.font.size = Pt(10)
    run.font.color.rgb = ACCENT
    # A bottom border reads as a rule to humans and is ignored by ATS parsers.
    p_pr = p._p.get_or_add_pPr()
    borders = p_pr.makeelement(
        "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}pBdr", {}
    )
    bottom = p_pr.makeelement(
        "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}bottom",
        {
            "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val": "single",
            "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}sz": "6",
            "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}space": "2",
            "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}color": "9AA3AF",
        },
    )
    borders.append(bottom)
    p_pr.append(borders)


def render_cv_docx(cv: GeneratedCV, path: Path) -> Path:
    doc = Document()
    _configure_styles(doc)

    name_p = doc.add_paragraph()
    name_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name_p.paragraph_format.space_after = Pt(2)
    name_run = name_p.add_run(cv.full_name or "")
    name_run.bold = True
    name_run.font.size = Pt(19)
    name_run.font.color.rgb = ACCENT

    if cv.headline:
        h = doc.add_paragraph()
        h.alignment = WD_ALIGN_PARAGRAPH.CENTER
        h.paragraph_format.space_after = Pt(2)
        hr = h.add_run(cv.headline)
        hr.font.size = Pt(11)
        hr.font.color.rgb = MUTED

    if cv.contact:
        c = doc.add_paragraph()
        c.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cr = c.add_run("  |  ".join(cv.contact))
        cr.font.size = Pt(9)
        cr.font.color.rgb = MUTED

    if cv.summary:
        _heading(doc, "Professional Summary")
        doc.add_paragraph(cv.summary)

    if cv.skills:
        _heading(doc, "Skills")
        # Comma separated rather than bullet glyphs: ATS keyword extraction is
        # more reliable on plain text, and it keeps the file fully ASCII.
        doc.add_paragraph(", ".join(cv.skills))

    if cv.experience:
        _heading(doc, "Experience")
        for exp in cv.experience:
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(7)
            p.paragraph_format.space_after = Pt(1)
            # Keep the role, its dates and the first bullet together, so a longer CV
            # cannot orphan a job heading at the foot of a page.
            p.paragraph_format.keep_with_next = True
            role = p.add_run(exp.role or "")
            role.bold = True
            role.font.size = Pt(11)
            if exp.company:
                # Comma, not an em dash: the dash is a visual machine tell and
                # some ATS parsers split role/company badly on it.
                sep = p.add_run(f", {exp.company}")
                sep.font.size = Pt(11)
                sep.font.color.rgb = ACCENT

            meta_bits = [b for b in (exp.dates, exp.location) if b]
            if meta_bits:
                m = doc.add_paragraph()
                m.paragraph_format.space_after = Pt(3)
                m.paragraph_format.keep_with_next = True
                mr = m.add_run("  |  ".join(meta_bits))
                mr.italic = True
                mr.font.size = Pt(9)
                mr.font.color.rgb = MUTED

            for bullet in exp.bullets:
                bp = doc.add_paragraph(bullet.text, style="List Bullet")
                bp.paragraph_format.space_after = Pt(2)

    if cv.projects:
        _heading(doc, "Projects")
        for proj in cv.projects:
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(5)
            p.paragraph_format.space_after = Pt(1)
            r = p.add_run(proj.name or "")
            r.bold = True
            if proj.technologies:
                t = p.add_run(f"  ({', '.join(proj.technologies)})")
                t.font.size = Pt(9)
                t.font.color.rgb = MUTED
            if proj.description:
                doc.add_paragraph(proj.description)

    if cv.education:
        _heading(doc, "Education")
        for edu in cv.education:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(1)
            r = p.add_run(edu.qualification or "")
            r.bold = True
            tail = ", ".join(b for b in (edu.institution, edu.year) if b)
            if tail:
                p.add_run(f", {tail}")

    path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(path)
    return path


def render_cover_letter_docx(text: str, cv: GeneratedCV, path: Path) -> Path:
    doc = Document()
    _configure_styles(doc)

    name_p = doc.add_paragraph()
    name_run = name_p.add_run(cv.full_name or "")
    name_run.bold = True
    name_run.font.size = Pt(15)
    name_run.font.color.rgb = ACCENT

    if cv.contact:
        c = doc.add_paragraph()
        c.paragraph_format.space_after = Pt(14)
        cr = c.add_run("  |  ".join(cv.contact))
        cr.font.size = Pt(9)
        cr.font.color.rgb = MUTED

    for block in (text or "").replace("\r\n", "\n").split("\n\n"):
        block = block.strip()
        if block:
            p = doc.add_paragraph(block)
            p.paragraph_format.space_after = Pt(9)

    path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(path)
    return path
